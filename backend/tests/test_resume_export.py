"""Tests for resume_export.py — pure logic, no external services."""

import os
import sys
import zipfile
from io import BytesIO

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from resume_export import (
    _parse_resume_sections,
    _strip_markdown,
    export_resume_docx,
    export_resume_pdf,
)

# Use the realistic resume from test_helpers
sys.path.insert(0, os.path.dirname(__file__))
from test_helpers import RESUME_TEXT


# ---------------------------------------------------------------------------
# _strip_markdown tests
# ---------------------------------------------------------------------------
class TestStripMarkdown:
    def test_removes_bold(self):
        assert _strip_markdown("**bold text**") == "bold text"

    def test_removes_italic(self):
        assert _strip_markdown("*italic*") == "italic"

    def test_removes_header_hashes(self):
        assert _strip_markdown("## Section Title") == "Section Title"

    def test_removes_dividers(self):
        assert _strip_markdown("---") == ""

    def test_mixed_markdown(self):
        result = _strip_markdown("## **Bold Header**")
        assert result == "Bold Header"

    def test_plain_text_unchanged(self):
        assert _strip_markdown("Just plain text") == "Just plain text"

    def test_empty_string(self):
        assert _strip_markdown("") == ""

    def test_nested_bold_italic(self):
        result = _strip_markdown("**some *nested* text**")
        assert "nested" in result
        assert "*" not in result


# ---------------------------------------------------------------------------
# _parse_resume_sections tests
# ---------------------------------------------------------------------------
class TestParseResumeSections:
    def test_header_lines_extracted(self):
        header, sections = _parse_resume_sections(RESUME_TEXT)
        assert len(header) >= 1
        assert "MIKE VOGT" in header[0]

    def test_sections_extracted(self):
        header, sections = _parse_resume_sections(RESUME_TEXT)
        section_titles = [t.lower() for t, _ in sections]
        # Our resume has SUMMARY, CORE COMPETENCIES, EXPERIENCE, EDUCATION
        assert any("summary" in t for t in section_titles)
        assert any("experience" in t for t in section_titles)
        assert any("education" in t for t in section_titles)

    def test_section_content_non_empty(self):
        header, sections = _parse_resume_sections(RESUME_TEXT)
        for title, content in sections:
            # At least one non-empty content line per section
            non_empty = [c for c in content if c.strip()]
            assert len(non_empty) >= 1, f"Section '{title}' has no content"

    def test_no_trailing_empty_lines(self):
        header, sections = _parse_resume_sections(RESUME_TEXT)
        for title, content in sections:
            if content:
                assert content[-1] != "", f"Section '{title}' has trailing empty line"

    def test_empty_input(self):
        header, sections = _parse_resume_sections("")
        assert header == []
        assert sections == []

    def test_no_sections_all_header(self):
        text = "John Doe\njohn@example.com\n555-1234"
        header, sections = _parse_resume_sections(text)
        assert len(header) == 3
        assert sections == []

    def test_case_insensitive_section_detection(self):
        text = "Name\n\nSUMMARY\nA great summary.\n\nskills\nPython, Java"
        header, sections = _parse_resume_sections(text)
        section_titles = [t for t, _ in sections]
        assert len(section_titles) >= 1

    def test_markdown_headers_detected(self):
        text = "Name\n\n## Summary\nA summary.\n\n## Experience\nDid things."
        header, sections = _parse_resume_sections(text)
        assert len(sections) == 2


# ---------------------------------------------------------------------------
# export_resume_pdf tests
# ---------------------------------------------------------------------------
class TestExportResumePdf:
    def test_returns_bytes(self):
        result = export_resume_pdf(RESUME_TEXT)
        assert isinstance(result, BytesIO)

    def test_pdf_magic_bytes(self):
        result = export_resume_pdf(RESUME_TEXT)
        data = result.read()
        assert data[:5] == b"%PDF-"

    def test_non_zero_size(self):
        result = export_resume_pdf(RESUME_TEXT)
        data = result.read()
        assert len(data) > 100

    def test_empty_input_produces_valid_pdf(self):
        result = export_resume_pdf("")
        data = result.read()
        assert data[:5] == b"%PDF-"

    def test_long_resume_multi_page(self):
        long_text = "Name\n\nExperience\n" + ("- Accomplished great things\n" * 200)
        result = export_resume_pdf(long_text)
        data = result.read()
        assert len(data) > 500


# ---------------------------------------------------------------------------
# export_resume_docx tests
# ---------------------------------------------------------------------------
class TestExportResumeDocx:
    def test_returns_bytes(self):
        result = export_resume_docx(RESUME_TEXT)
        assert isinstance(result, BytesIO)

    def test_valid_docx_is_zip(self):
        """DOCX files are ZIP archives."""
        result = export_resume_docx(RESUME_TEXT)
        assert zipfile.is_zipfile(result)

    def test_non_zero_size(self):
        result = export_resume_docx(RESUME_TEXT)
        data = result.read()
        assert len(data) > 100

    def test_empty_input_produces_valid_docx(self):
        result = export_resume_docx("")
        assert zipfile.is_zipfile(result)

    def test_docx_contains_content_xml(self):
        result = export_resume_docx(RESUME_TEXT)
        with zipfile.ZipFile(result) as zf:
            names = zf.namelist()
            assert "word/document.xml" in names

    def test_long_content_no_truncation(self):
        long_text = "Name\n\nExperience\n" + ("Detailed experience line.\n" * 300)
        result = export_resume_docx(long_text)
        data = result.read()
        assert len(data) > 1000


# ---------------------------------------------------------------------------
# PDF content verification tests (Phase 9 — Gap #8)
# ---------------------------------------------------------------------------
class TestPdfContentVerification:
    """Verify PDF text contains section headings, skills, and employer names."""

    def _extract_pdf_text(self, pdf_buf):
        """Extract text from PDF BytesIO buffer using pypdf."""
        from pypdf import PdfReader

        pdf_buf.seek(0)
        reader = PdfReader(pdf_buf)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def test_pdf_contains_section_headings(self):
        pdf_buf = export_resume_pdf(RESUME_TEXT)
        text = self._extract_pdf_text(pdf_buf)
        text_upper = text.upper()
        assert "SUMMARY" in text_upper, "PDF should contain SUMMARY heading"
        assert "EXPERIENCE" in text_upper, "PDF should contain EXPERIENCE heading"
        assert "EDUCATION" in text_upper, "PDF should contain EDUCATION heading"

    def test_pdf_contains_key_skills(self):
        pdf_buf = export_resume_pdf(RESUME_TEXT)
        text = self._extract_pdf_text(pdf_buf)
        text_lower = text.lower()
        assert "python" in text_lower, "PDF should mention Python"
        assert "docker" in text_lower, "PDF should mention Docker"

    def test_pdf_page_count_reasonable(self):
        from pypdf import PdfReader

        pdf_buf = export_resume_pdf(RESUME_TEXT)
        pdf_buf.seek(0)
        reader = PdfReader(pdf_buf)
        assert (
            1 <= len(reader.pages) <= 3
        ), f"PDF has {len(reader.pages)} pages — expected 1-3 for a resume"

    def test_pdf_contains_employer_names(self):
        pdf_buf = export_resume_pdf(RESUME_TEXT)
        text = self._extract_pdf_text(pdf_buf)
        assert (
            "Navitus" in text or "OPI" in text or "AHEAD" in text
        ), "PDF should contain at least one employer name"


# ---------------------------------------------------------------------------
# DOCX content verification tests (Phase 9 — Gap #8)
# ---------------------------------------------------------------------------
class TestDocxContentVerification:
    """Verify DOCX contains section headings, skills, bullet paragraphs."""

    def _read_docx_text(self, docx_buf):
        """Read full text from DOCX BytesIO buffer."""
        from docx import Document

        docx_buf.seek(0)
        doc = Document(docx_buf)
        return "\n".join(p.text for p in doc.paragraphs)

    def _get_docx_paragraphs(self, docx_buf):
        from docx import Document

        docx_buf.seek(0)
        return Document(docx_buf).paragraphs

    def test_docx_contains_section_headings(self):
        docx_buf = export_resume_docx(RESUME_TEXT)
        text = self._read_docx_text(docx_buf)
        text_upper = text.upper()
        assert "SUMMARY" in text_upper, "DOCX should contain SUMMARY"
        assert "EXPERIENCE" in text_upper, "DOCX should contain EXPERIENCE"
        assert "EDUCATION" in text_upper, "DOCX should contain EDUCATION"

    def test_docx_contains_key_skills(self):
        docx_buf = export_resume_docx(RESUME_TEXT)
        text = self._read_docx_text(docx_buf)
        text_lower = text.lower()
        assert "python" in text_lower, "DOCX should mention Python"
        assert "docker" in text_lower, "DOCX should mention Docker"

    def test_docx_has_multiple_paragraphs(self):
        """DOCX has multiple content paragraphs for a realistic resume."""
        docx_buf = export_resume_docx(RESUME_TEXT)
        paragraphs = self._get_docx_paragraphs(docx_buf)
        non_empty = [p for p in paragraphs if p.text.strip()]
        assert len(non_empty) >= 5, f"Expected ≥5 non-empty paragraphs, found {len(non_empty)}"

    def test_docx_name_is_present(self):
        """First paragraph contains the name."""
        docx_buf = export_resume_docx(RESUME_TEXT)
        paragraphs = self._get_docx_paragraphs(docx_buf)
        # Find the first non-empty paragraph
        first_text = ""
        for p in paragraphs:
            if p.text.strip():
                first_text = p.text
                break
        assert (
            "MIKE VOGT" in first_text.upper()
        ), f"First paragraph should contain name, got: {first_text!r}"

    def test_optimized_pdf_contains_added_keywords(self):
        """Optimize → export PDF → text includes optimization content."""
        from nlp_engine import extract_keywords
        from pypdf import PdfReader
        from utils import optimize_resume

        resume_data = {"text": RESUME_TEXT, "skills": extract_keywords(RESUME_TEXT, 20)}
        job_keywords = extract_keywords(
            "We need Python, Docker, Kubernetes, Terraform, REST APIs, CI/CD, PostgreSQL",
            20,
        )
        result = optimize_resume(
            resume_data,
            job_keywords,
            job_text="We need Python, Docker, Kubernetes, Terraform, REST APIs, CI/CD, PostgreSQL",
        )
        pdf_buf = export_resume_pdf(result["optimized_text"])
        pdf_buf.seek(0)
        reader = PdfReader(pdf_buf)
        pdf_text = ""
        for page in reader.pages:
            pdf_text += page.extract_text() or ""
        assert "python" in pdf_text.lower(), "Optimized PDF should mention Python"

    def test_optimized_docx_preserves_sections(self):
        """Optimize → export DOCX → paragraphs include section content."""
        from nlp_engine import extract_keywords
        from utils import optimize_resume

        resume_data = {"text": RESUME_TEXT, "skills": extract_keywords(RESUME_TEXT, 20)}
        job_keywords = extract_keywords(
            "We need Python, Docker, Kubernetes, Terraform, REST APIs",
            20,
        )
        result = optimize_resume(
            resume_data,
            job_keywords,
            job_text="We need Python, Docker, Kubernetes, Terraform, REST APIs",
        )
        docx_buf = export_resume_docx(result["optimized_text"])
        text = self._read_docx_text(docx_buf)
        assert len(text) > 100, "Optimized DOCX should have substantial content"
